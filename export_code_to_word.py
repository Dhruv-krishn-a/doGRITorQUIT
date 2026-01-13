#!/usr/bin/env python3
"""
Export all source code files in the current project to a single Word document.
Each file is written as:

// relative/path/to/file.ext
<file contents>
"""

import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ================= CONFIG =================

OUTPUT_FILE = "project_code_dump.docx"

CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".tsx", ".jsx",
    ".java", ".cpp", ".c", ".h",
    ".cs", ".go", ".rs",
    ".html", ".htm", ".css", ".scss",
    ".json", ".yml", ".yaml",
    ".sh", ".bash", ".md", ".xml"
}

# ✅ ALL directories to ignore
IGNORE_DIRS = {
    "node_modules",
    ".git",
    ".next",
    "dist",
    "build",
    "out",
    ".expo",
    ".expo-shared",
    "__pycache__",
    ".turbo",
    ".vercel",
    ".cache"
}

MAX_FILE_SIZE_BYTES = 5 * 1024 * 1024  # 5 MB

# =========================================

CONTROL_CHAR_RE = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')


def is_code_file(filename: str) -> bool:
    return any(filename.lower().endswith(ext) for ext in CODE_EXTENSIONS)


def collect_code_files(root_dir: str):
    files = []
    for root, dirs, filenames in os.walk(root_dir):
        # 🔥 THIS is the critical ignore logic
        dirs[:] = [d for d in dirs if d not in IGNORE_DIRS]

        for filename in filenames:
            if is_code_file(filename):
                files.append(os.path.join(root, filename))

    return sorted(files)


def clean_text(text: str) -> str:
    """Remove characters that break python-docx / Word XML."""
    return CONTROL_CHAR_RE.sub("", text)


def add_code_block(doc: Document, rel_path: str, code: str):
    # File path header
    header = doc.add_paragraph()
    run = header.add_run(f"// {rel_path}")
    run.bold = True

    # Code paragraph (monospace)
    para = doc.add_paragraph()
    for line in code.splitlines():
        r = para.add_run(line)
        try:
            r.font.name = "Courier New"
            r.font.size = Pt(9)
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        except Exception:
            pass
        r.add_break()


def main():
    root_dir = os.getcwd()
    print(f"📂 Scanning project: {root_dir}")

    doc = Document()
    doc.add_heading("Project Code Dump", level=1)

    files = collect_code_files(root_dir)
    print(f"📄 Found {len(files)} code files")

    for file_path in files:
        rel_path = os.path.relpath(file_path, root_dir)

        try:
            if os.path.getsize(file_path) > MAX_FILE_SIZE_BYTES:
                print(f"⏭ Skipped (too large): {rel_path}")
                continue

            with open(file_path, "rb") as f:
                raw = f.read()

            try:
                text = raw.decode("utf-8")
            except UnicodeDecodeError:
                text = raw.decode("utf-8", errors="replace")

            text = clean_text(text)
            add_code_block(doc, rel_path, text)
            print(f"✅ Added: {rel_path}")

        except Exception as e:
            print(f"❌ Error reading {rel_path}: {e}")

    doc.save(OUTPUT_FILE)
    print(f"\n🎉 Done! Output saved as: {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
